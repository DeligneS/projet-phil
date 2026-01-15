"""Word document parser using python-docx."""

import io

from docx import Document


def parse_docx(content: bytes) -> str:
    """Extract text content from a Word document.

    Args:
        content: Raw bytes of the DOCX file.

    Returns:
        Extracted text from all paragraphs and tables.
    """
    doc = Document(io.BytesIO(content))
    text_parts = []

    # Extract text from paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n\n".join(text_parts)
