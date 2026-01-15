"""Word document export for evaluation results."""

import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.evaluation.llm_evaluator import EvaluationResult


def create_word_document(evaluation: EvaluationResult) -> io.BytesIO:
    """Create a Word document for a single student evaluation.

    Args:
        evaluation: Evaluation result for a student.

    Returns:
        BytesIO containing the Word document.
    """
    doc = Document()

    # Title
    title = doc.add_heading(f"Évaluation - {evaluation.student_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Final grade section
    doc.add_heading("Note finale", level=1)
    grade_para = doc.add_paragraph()
    grade_run = grade_para.add_run(f"{evaluation.note_finale} / {evaluation.note_max}")
    grade_run.bold = True
    grade_run.font.size = Pt(16)

    # Criteria section
    doc.add_heading("Évaluation par critère", level=1)

    for criterion in evaluation.criteres:
        # Criterion name as heading
        doc.add_heading(criterion.nom, level=2)

        # Score
        score_para = doc.add_paragraph()
        score_para.add_run("Note : ").bold = True
        score_para.add_run(f"{criterion.note} / {criterion.note_max}")

        # Comment
        comment_para = doc.add_paragraph()
        comment_para.add_run("Commentaire : ").bold = True
        comment_para.add_run(criterion.commentaire)

    # General feedback section
    doc.add_heading("Feedback général", level=1)
    doc.add_paragraph(evaluation.feedback_general)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def create_word_documents_zip(evaluations: list[EvaluationResult]) -> io.BytesIO:
    """Create a ZIP containing Word documents for all students.

    Args:
        evaluations: List of evaluation results.

    Returns:
        BytesIO containing the ZIP file.
    """
    import zipfile

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for evaluation in evaluations:
            # Create Word document
            doc_buffer = create_word_document(evaluation)

            # Sanitize filename
            safe_name = "".join(
                c if c.isalnum() or c in (" ", "-", "_") else "_"
                for c in evaluation.student_name
            )
            filename = f"{safe_name}.docx"

            # Add to ZIP
            zf.writestr(filename, doc_buffer.getvalue())

    zip_buffer.seek(0)
    return zip_buffer


def create_free_format_word_document(
    student_name: str,
    content: str,
) -> io.BytesIO:
    """Create a Word document with free-form content.

    Args:
        student_name: Name of the student.
        content: The text content to include.

    Returns:
        BytesIO containing the Word document.
    """
    doc = Document()

    # Title
    title = doc.add_heading(f"Évaluation - {student_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Content - split by double newlines for paragraphs
    paragraphs = content.split("\n\n")
    for para_text in paragraphs:
        if para_text.strip():
            # Check if it looks like a heading (starts with # or is short and uppercase)
            stripped = para_text.strip()
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            else:
                doc.add_paragraph(stripped)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def create_free_format_documents_zip(
    evaluations: list[tuple[str, str]],
) -> io.BytesIO:
    """Create a ZIP containing free-format Word documents.

    Args:
        evaluations: List of (student_name, content) tuples.

    Returns:
        BytesIO containing the ZIP file.
    """
    import zipfile

    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for student_name, content in evaluations:
            # Create Word document
            doc_buffer = create_free_format_word_document(student_name, content)

            # Sanitize filename
            safe_name = "".join(
                c if c.isalnum() or c in (" ", "-", "_") else "_"
                for c in student_name
            )
            filename = f"{safe_name}.docx"

            # Add to ZIP
            zf.writestr(filename, doc_buffer.getvalue())

    zip_buffer.seek(0)
    return zip_buffer
